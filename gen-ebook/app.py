import streamlit as st
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain.chains import LLMChain
from langchain.chains import SequentialChain
import json
from docx import Document
#from reportlab.lib.pagesizes import letter
#from reportlab.pdfgen import canvas

# Load environment variables
load_dotenv()

# Initialize language model

llms = OpenAI(temperature=0.6, verbose=True, max_tokens=3500)

# Initialize Prompt Templates
prompt_template_title = PromptTemplate(
    input_variables=['topic'],
    template="""You are an AI Assistant who is helping an author write an ebook about book about {topic}. I want you to generate a good title for such a book.  
    The title should be between 6 to 10 words.  The title should be in title case."""
)

title_chain = LLMChain(llm=llms, prompt=prompt_template_title, output_key="title", verbose=True)

prompt_template_chapters = PromptTemplate(
    input_variables=['title', 'topic'],
    template="""
    I want you to act as an expert on the topic of '{topic}'. I want you to create a complete outline for a book called '{title}' for that topic'.  
    The outline should be as detailed as possible.  
    The book should have at least 8 chapters and each chapter should have 3-5 sub topics. These subtopics should becompletely defined as well. 
    The output should be in JSON format.  Only return the JSON structure.
    the JSON structure should be a list of chapters.  Example JSON:
{{
   "chapters":[
      {{
         "title":"Introduction to Agile Methodologies",
         "subTopics":[
            {{
               "title":"What is Agile?",
               "subTopics":[
                  "History of Agile",
                  "Agile Principles and Values",
                  "The Agile Manifesto"
               ]
            }},
            {{
               "title":"Benefits of Agile Methodologies",
               "subTopics":[
                  "Increased Productivity",
                  "Better Quality",
                  "Faster Delivery",
                  "Improved Employee Engagement"
               ]
            }}
         ]
      }},
      {{
         "title":"Agile Methodologies and Practices",
         "subTopics":[
            {{
               "title":"Scrum",
               "subTopics":[
                  "Roles and Responsibilities",
                  "Scrum Events",
                  "Sprint Planning and Retrospective"
               ]
            }}
         ]
      }}
   ]
}}
    """
)

chapter_chain = LLMChain(llm=llms, prompt=prompt_template_chapters, output_key="book_chapters", verbose=True)

# New Prompt Template for Chapter Details
prompt_template_chapter_details = PromptTemplate(
    input_variables=['chapter_title', 'topic'],
    template="Please generate the detail text for a chapter in a book about '{topic}' entitled '{chapter_title}'.  Be sure to include as much detail as possible"
)

# New LLMChain for Chapter Details
chapter_details_chain = LLMChain(llm=llms, prompt=prompt_template_chapter_details, output_key="chapter_details", verbose=True)

# New Prompt Template for Subtopic Details
prompt_template_subtopic_details = PromptTemplate(
    input_variables=['subtopic_title', 'topic', 'style'],
    template="""Please generate the detail text for a subtopic in a book about '{topic}' entitled '{subtopic_title}'. 
    Be sure to include as much detail as possible. It should be written in the writing style similar to '{style}'"""
)

# New LLMChain for Subtopic Details
subtopic_details_chain = LLMChain(llm=llms, prompt=prompt_template_subtopic_details, output_key="subtopic_details", verbose=True)


# Initialize the SequentialChain
chain = SequentialChain(
    chains=[title_chain, chapter_chain],
    input_variables=["topic"],
    output_variables=["title", "book_chapters"]
)

# Streamlit UI
st.title('Generate Book')

# Display a warning message about the potential cost
st.warning("""
⚠️ Generating a full ebook will result in a large number of requests to the language model 
and may incur substantial costs. 
Please proceed with caution.
""")

# Checkbox to confirm the user wants to proceed
continue_box = st.checkbox("I understand and want to continue.")

# If the user doesn't check the box, don't display the rest of the UI.
if not continue_box:
    st.stop()


# Text input for the topic
topic = st.text_input('Enter the topic for the book:', '')

# Text input for the style
style = st.text_input('Author Style for the book:', 'Ken Burns')


# Text input for the file name
file_name = st.text_input('Enter the name for the output file:', '')

# Ensure the file name ends with '.docx'
if not file_name.lower().endswith('.docx'):
    file_name += '.docx'

# Selection for the output format
output_format = st.selectbox('Choose the output format:', ['Word'])

# Initialize new Document
doc = Document()

# Button to generate the title and chapters
if st.button('Generate'):
    if topic and file_name:
        response = chain({'topic': topic, 'style': style})
        # Add title to Document
        doc.add_heading(f"{response['title']}", 0)
 
        st.write(f"{response['title']}")
        
        parsed_chapters = json.loads(response['book_chapters']) if isinstance(response['book_chapters'], str) else response['book_chapters']
        
        #Generate a Table of Contents
        doc.add_heading('Table of Contents', level=1)
        toc = doc.add_paragraph()
        
        for i, chapter in enumerate(parsed_chapters['chapters']):
            st.write(f"Chapter {i + 1} - {chapter['title']}")
            toc.add_run(f"{i + 1} - {chapter['title']}").bold = True
            toc.add_run('\n')
            for j, subtopic in enumerate(chapter['subTopics']):
                st.write(f"\t{j + 1} - {subtopic['title']}")
                toc.add_run(f"\t{j + 1} - {subtopic['title']}")
                toc.add_run('\n')

        # Add a page break

        doc.add_page_break()

        
        for i, chapter in enumerate(parsed_chapters['chapters']):
             # Add chapter title to Document
            doc.add_heading(f"Chapter {i + 1} - {chapter['title']}", level=1)
         
            st.write(f"Chapter {i + 1} - {chapter['title']}")
            
            # Fetch and display chapter details
            chapter_details_response = chapter_details_chain({'chapter_title': chapter['title'], 'topic': topic})

            # Add chapter details to Document
            doc.add_paragraph(chapter_details_response['chapter_details'])
          
            st.write(chapter_details_response['chapter_details'])
            
            for j, subtopic in enumerate(chapter['subTopics']):
                # Add subtopic title to Document
                doc.add_heading(f"{subtopic['title']}", level=2)
          
                st.write(f"{subtopic['title']}")
                
                # Fetch and display subtopic details
                subtopic_details_response = subtopic_details_chain({'subtopic_title': subtopic['title'], 'topic': topic, 'style': style})

                # Add subtopic details to Document
                doc.add_paragraph(subtopic_details_response['subtopic_details'])
                st.write(subtopic_details_response['subtopic_details'])
                
            st.write("")  # Adds a line break for better readability
            break
        # Save the document
        doc.save(file_name)
        st.write(f"The generated book has been saved as '{file_name}'")

    else:
        st.write("Please enter a topic and file name.")